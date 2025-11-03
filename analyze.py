#!/usr/bin/env python3
"""
Writing Style Analyzer - Analyze writing style using local LLMs
Supports German and English text analysis
"""

import argparse
import json
import logging
import re
import sys
from collections import Counter
from datetime import datetime
from pathlib import Path
from typing import Any

import pdfplumber
import yaml
from docx import Document
from langdetect import detect_langs
from odf import teletype, text
from odf.opendocument import load as load_odt
from tqdm import tqdm


class ConfigurationError(Exception):
    """Raised when configuration is invalid"""

    pass


class AnalysisError(Exception):
    """Raised when analysis fails"""

    pass


def setup_logging(config: dict[str, Any]) -> logging.Logger:
    """Setup logging with config parameters"""
    log_config = config.get("logging", {})
    level = getattr(logging, log_config.get("level", "INFO"))

    # Create logger
    logger = logging.getLogger("wsa")
    logger.setLevel(level)

    # Console handler
    console_handler = logging.StreamHandler()
    console_handler.setLevel(level)
    console_formatter = logging.Formatter("%(asctime)s - %(name)s - %(levelname)s - %(message)s")
    console_handler.setFormatter(console_formatter)
    logger.addHandler(console_handler)

    # File handler if configured
    log_file = log_config.get("file")
    if log_file:
        file_handler = logging.FileHandler(log_file)
        file_handler.setLevel(level)
        file_handler.setFormatter(console_formatter)
        logger.addHandler(file_handler)

    return logger


def load_config(config_path: str = "config.yaml") -> dict[str, Any]:
    """Load configuration from YAML file"""
    config_file = Path(config_path)
    if not config_file.exists():
        raise ConfigurationError(f"Configuration file not found: {config_path}")

    with open(config_file, encoding="utf-8") as f:
        config = yaml.safe_load(f)

    return config


def load_linguistic_patterns(language: str = "de", style: str = "academic") -> dict[str, Any]:
    """Load base linguistic patterns from YAML files

    Args:
        language: Language code (de, en)
        style: Writing style (academic, conversational, etc.)

    Returns:
        Dictionary with transition word patterns by category
    """
    patterns_dir = Path(__file__).parent / "linguistic_patterns" / "base"
    pattern_file = patterns_dir / f"{language}_{style}.yaml"

    if not pattern_file.exists():
        # Fallback to empty patterns if file doesn't exist
        return {}

    try:
        with open(pattern_file, encoding="utf-8") as f:
            data = yaml.safe_load(f)
            return data.get("transitions", {})
    except Exception as e:
        logging.warning(f"Failed to load linguistic patterns from {pattern_file}: {e}")
        return {}


def discover_transition_words_llm(
    text: str, language: str, llm_analyzer, max_tokens: int = 500
) -> dict[str, list[str]]:
    """Use LLM to discover transition words in the given text

    This function asks the LLM to identify transition/connector words that appear
    in the text, categorized by function. This captures user-specific patterns
    beyond the base patterns.

    Args:
        text: Text sample to analyze (should be representative)
        language: Language code (de, en)
        llm_analyzer: LLMAnalyzer instance
        max_tokens: Max tokens for LLM response

    Returns:
        Dictionary mapping categories to lists of discovered transition words
    """
    # Use strategic sampling to get representative text (beginning + middle + end)
    text_sample = llm_analyzer._sample_text_strategically(text, max_length=6000)

    if language == "de":
        prompt = f"""Analysiere den folgenden deutschen Text und identifiziere ALLE Übergangswörter/Konnektoren, die tatsächlich im Text vorkommen.

Kategorisiere sie nach Funktion:
- additive (hinzufügen): außerdem, zudem, etc.
- contrastive (gegenüberstellen): jedoch, allerdings, etc.
- causal (Ursache/Wirkung): daher, deshalb, etc.
- temporal (zeitlich): zunächst, dann, etc.
- conclusive (zusammenfassend): insgesamt, abschließend, etc.

Gib NUR die Wörter zurück, die TATSÄCHLICH im Text erscheinen.

Text:
{text_sample}

Antworte im JSON-Format:
{{"additive": ["wort1", "wort2"], "contrastive": [...], ...}}"""
    else:  # English
        prompt = f"""Analyze the following English text and identify ALL transition words/connectors that actually appear in the text.

Categorize them by function:
- additive (adding): furthermore, moreover, etc.
- contrastive (contrasting): however, nevertheless, etc.
- causal (cause/effect): therefore, thus, etc.
- temporal (time): first, then, etc.
- conclusive (summarizing): in conclusion, overall, etc.

Return ONLY words that ACTUALLY appear in the text.

Text:
{text_sample}

Respond in JSON format:
{{"additive": ["word1", "word2"], "contrastive": [...], ...}}"""

    try:
        response = llm_analyzer.generate(prompt, max_tokens=max_tokens)
        # Extract JSON from response
        discovered = llm_analyzer._extract_json_from_response(response)
        if discovered and isinstance(discovered, dict):
            return discovered
        return {}
    except Exception as e:
        logging.warning(f"LLM pattern discovery failed: {e}")
        return {}


class TextProcessor:
    """Process and analyze text structure"""

    def __init__(self, logger: logging.Logger):
        self.logger = logger

    def detect_language(self, text: str) -> tuple[str, float]:
        """Detect the primary language of text with confidence score"""
        try:
            langs = detect_langs(text)
            if langs:
                primary = langs[0]
                return primary.lang, primary.prob
            return "unknown", 0.0
        except Exception as e:
            self.logger.warning(f"Language detection failed: {e}")
            return "unknown", 0.0

    def split_sentences(self, text: str) -> list[str]:
        """Split text into sentences (works for German and English)"""
        # Handle common abbreviations
        text = re.sub(r"\b(Dr|Prof|etc|z\.B|d\.h|bzw)\.", r"\1<DOT>", text)

        # Split on sentence endings
        sentences = re.split(r"[.!?]+\s+", text)

        # Restore abbreviations
        sentences = [s.replace("<DOT>", ".") for s in sentences if s.strip()]

        return sentences

    def split_paragraphs(self, text: str) -> list[str]:
        """Split text into paragraphs"""
        # Split on multiple newlines
        paragraphs = re.split(r"\n\s*\n", text)
        return [p.strip() for p in paragraphs if p.strip()]

    def tokenize_words(self, text: str) -> list[str]:
        """Tokenize text into words (basic whitespace tokenization)"""
        # Remove punctuation and split
        words = re.findall(r"\b\w+\b", text.lower())
        return words

    def is_code_line(self, line: str) -> bool:
        """Detect if a line looks like code rather than prose"""
        line = line.strip()
        if not line:
            return False

        # Code indicators
        code_patterns = [
            r"^\s*(function|end|if|else|for|while|return|var|let|const|def|class)\b",  # Keywords
            r"^\s*[a-zA-Z_]\w*\s*[=\(\[]",  # Variable assignments or function calls
            r"[{};]\s*$",  # Code block markers
            r"^\s*%",  # MATLAB comments
            r"^\s*//",  # C-style comments
            r"^\s*#",  # Python comments (but not markdown headers with text)
            r"^\s*\d+\s*$",  # Line numbers only
            r"[=+\-*/]{2,}",  # Multiple operators
            r"\)\s*;",  # Function call endings
        ]

        for pattern in code_patterns:
            if re.search(pattern, line):
                return True

        # High density of special characters suggests code
        special_chars = len(re.findall(r"[{}()\[\];=<>]", line))
        if len(line) > 0 and special_chars / len(line) > 0.2:
            return True

        return False

    def is_reference_line(self, line: str) -> bool:
        """Detect if a line is a bibliographic reference"""
        line = line.strip()
        if not line:
            return False

        # Reference indicators
        reference_patterns = [
            r"^\[\d+\]",  # [1] citation
            r"^https?://",  # URLs
            r"doi:",  # DOI identifiers
            r"^\d+\.\s+[A-Z].*\(\d{4}\)",  # "1. Author Name (2024)"
            r"et al\.",  # "et al." citation marker
            r"pp\.\s*\d+",  # Page numbers
            r"Vol\.\s*\d+",  # Volume numbers
        ]

        for pattern in reference_patterns:
            if re.search(pattern, line, re.IGNORECASE):
                return True

        return False

    def is_likely_prose(self, text: str) -> tuple[bool, float]:
        """Determine if text is likely prose (not code/tables/references)"""
        lines = text.split("\n")
        if not lines:
            return False, 0.0

        code_lines = 0
        ref_lines = 0
        prose_lines = 0

        for line in lines:
            line = line.strip()
            if not line:
                continue

            if self.is_code_line(line):
                code_lines += 1
            elif self.is_reference_line(line):
                ref_lines += 1
            else:
                prose_lines += 1

        total_lines = code_lines + ref_lines + prose_lines
        if total_lines == 0:
            return False, 0.0

        prose_ratio = prose_lines / total_lines
        is_prose = prose_ratio > 0.6  # At least 60% prose
        return is_prose, prose_ratio

    def filter_to_prose(self, text: str, min_words_per_paragraph: int = 10) -> str:
        """Filter text to keep only prose paragraphs

        Args:
            text: Input text
            min_words_per_paragraph: Minimum words required (default 10, was 20 before)
        """
        paragraphs = self.split_paragraphs(text)
        prose_paragraphs = []

        for para in paragraphs:
            # Skip very short paragraphs (likely headers, fragments)
            word_count = len(para.split())
            if word_count < 5:  # Too short to be meaningful
                continue

            is_prose, prose_ratio = self.is_likely_prose(para)

            # More lenient: accept if >40% prose (was 60%) OR long enough
            if (prose_ratio > 0.4 and word_count >= min_words_per_paragraph) or word_count >= 30:
                # Additional filtering: remove lines that are clearly code/references
                filtered_lines = []
                for line in para.split("\n"):
                    if not self.is_code_line(line) and not self.is_reference_line(line):
                        filtered_lines.append(line)

                filtered_para = "\n".join(filtered_lines).strip()
                if filtered_para and len(filtered_para.split()) >= 5:  # Still has content
                    prose_paragraphs.append(filtered_para)

        return "\n\n".join(prose_paragraphs)

    def calculate_lexical_diversity(self, words: list[str]) -> float:
        """Calculate lexical diversity (unique words / total words)"""
        if not words:
            return 0.0
        unique_words = len(set(words))
        total_words = len(words)
        return unique_words / total_words if total_words > 0 else 0.0

    def detect_german_features(self, text: str, words: list[str]) -> dict[str, Any]:
        """Detect German-specific linguistic features"""
        features = {}

        # Detect compound words (rough heuristic: long words)
        long_words = [w for w in words if len(w) > 15]
        features["has_compound_words"] = len(long_words) > 0
        features["compound_word_examples"] = long_words[:5] if long_words else []

        # Detect formality (du vs Sie)
        du_pattern = r"\b(du|dich|dir|dein|deine)\b"
        sie_pattern = r"\b(Sie|Ihnen|Ihr|Ihre)\b"

        du_count = len(re.findall(du_pattern, text, re.IGNORECASE))
        sie_count = len(re.findall(sie_pattern, text, re.IGNORECASE))

        if du_count > sie_count:
            features["formality"] = "informal (du-form)"
        elif sie_count > du_count:
            features["formality"] = "formal (Sie-form)"
        else:
            features["formality"] = "neutral"

        # Detect umlauts
        umlaut_pattern = r"[äöüÄÖÜß]"
        features["uses_umlauts"] = bool(re.search(umlaut_pattern, text))

        return features

    def analyze_voice(self, text: str, language: str = "de") -> dict[str, Any]:
        """Analyze passive vs active voice usage"""
        sentences = self.split_sentences(text)

        if language == "de":
            # German passive indicators
            passive_patterns = [
                r"\bwurde(n)?\b",  # wurde, wurden
                r"\bwird\b",  # wird
                r"\bwerden\b",  # werden
                r"\bworden\b",  # worden
                r"\bgew\w+\b",  # geworden, gewesen, etc.
            ]
        else:  # English
            passive_patterns = [
                r"\b(is|are|was|were|been|being)\s+\w+(ed|en)\b",
                r"\bgets?\s+\w+(ed|en)\b",
            ]

        passive_count = 0
        for sentence in sentences:
            for pattern in passive_patterns:
                if re.search(pattern, sentence, re.IGNORECASE):
                    passive_count += 1
                    break

        total_sentences = len(sentences)
        passive_ratio = passive_count / total_sentences if total_sentences > 0 else 0.0

        return {
            "passive_sentences": passive_count,
            "total_sentences": total_sentences,
            "passive_ratio": round(passive_ratio, 3),
            "voice_style": (
                "mostly passive"
                if passive_ratio > 0.5
                else "balanced" if passive_ratio > 0.2 else "mostly active"
            ),
        }

    def extract_transition_words(
        self,
        text: str,
        language: str = "de",
        base_patterns: dict[str, Any] | None = None,
        discovered_patterns: dict[str, list[str]] | None = None,
    ) -> dict[str, Any]:
        """Extract and count transition/connector words using hybrid pattern system

        Args:
            text: Text to analyze
            language: Language code (de, en)
            base_patterns: Base patterns from YAML files (optional)
            discovered_patterns: LLM-discovered patterns from text (optional)

        Returns:
            Dictionary with transition word statistics
        """
        # Build transition word list from base patterns
        transitions = {}
        if base_patterns:
            # Extract word lists from base patterns
            for category, data in base_patterns.items():
                if isinstance(data, dict) and "words" in data:
                    transitions[category] = data["words"]
                elif isinstance(data, list):
                    transitions[category] = data
        else:
            # Fallback to empty dict - will be populated by discovered patterns
            transitions = {}

        # Merge with discovered patterns
        if discovered_patterns:
            for category, words in discovered_patterns.items():
                if category in transitions:
                    # Add discovered words to existing category (avoid duplicates)
                    existing = set(w.lower() for w in transitions[category])
                    for word in words:
                        if word.lower() not in existing:
                            transitions[category].append(word)
                else:
                    # New category from discovered patterns
                    transitions[category] = words

        text_lower = text.lower()
        found_transitions = {}
        total_transitions = 0

        for category, words in transitions.items():
            count = 0
            examples = []
            for word in words:
                word_count = len(re.findall(r"\b" + re.escape(word) + r"\b", text_lower))
                count += word_count
                if word_count > 0:
                    examples.append(word)

            if count > 0:
                found_transitions[category] = {"count": count, "examples": examples[:5]}
                total_transitions += count

        return {
            "total_transitions": total_transitions,
            "by_category": found_transitions,
            "transition_density": round(
                total_transitions / len(self.tokenize_words(text)) * 100, 2
            ),  # per 100 words
        }

    def analyze_sentence_complexity(self, text: str, language: str = "de") -> dict[str, Any]:
        """Analyze sentence structure complexity"""
        sentences = self.split_sentences(text)

        if language == "de":
            subordinate_patterns = [
                r"\b(weil|da|obwohl|wenn|als|nachdem|bevor|während|dass)\b",
            ]
        else:  # English
            subordinate_patterns = [
                r"\b(because|since|although|while|when|after|before|that|which|who)\b",
            ]

        complex_sentences = 0
        total_commas = 0
        total_semicolons = 0

        for sentence in sentences:
            # Count subordinating conjunctions
            for pattern in subordinate_patterns:
                if re.search(pattern, sentence, re.IGNORECASE):
                    complex_sentences += 1
                    break

            total_commas += sentence.count(",")
            total_semicolons += sentence.count(";")

        total_sentences = len(sentences)

        return {
            "complex_sentences": complex_sentences,
            "complexity_ratio": (
                round(complex_sentences / total_sentences, 3) if total_sentences > 0 else 0.0
            ),
            "avg_commas_per_sentence": (
                round(total_commas / total_sentences, 2) if total_sentences > 0 else 0.0
            ),
            "semicolons": total_semicolons,
            "structure_style": (
                (
                    "complex"
                    if complex_sentences / total_sentences > 0.5
                    else "mixed" if complex_sentences / total_sentences > 0.2 else "simple"
                )
                if total_sentences > 0
                else "unknown"
            ),
        }

    def analyze_rhetorical_devices(self, text: str) -> dict[str, Any]:
        """Analyze usage of rhetorical devices"""
        sentences = self.split_sentences(text)

        # Count questions
        questions = [s for s in sentences if "?" in s]

        # Count imperatives (rough heuristic: sentences starting with verbs)
        imperative_patterns = [
            r"^(Beachte|Nutze|Verwende|Probiere|Achte|Denke|Schau|Lies)",  # German
            r"^(Consider|Use|Try|Note|Think|Look|Read|Take)",  # English
        ]
        imperatives = 0
        for sentence in sentences:
            for pattern in imperative_patterns:
                if re.search(pattern, sentence.strip(), re.IGNORECASE):
                    imperatives += 1
                    break

        # Count lists (lines starting with -, *, numbers)
        list_items = len(re.findall(r"^\s*[-*•]\s", text, re.MULTILINE))
        list_items += len(re.findall(r"^\s*\d+\.\s", text, re.MULTILINE))

        # Count parenthetical remarks
        parentheses = text.count("(")
        dashes_pairs = text.count(" - ")

        return {
            "questions": len(questions),
            "question_ratio": (
                round(len(questions) / len(sentences), 3) if len(sentences) > 0 else 0.0
            ),
            "imperatives": imperatives,
            "list_items": list_items,
            "parenthetical_remarks": parentheses,
            "em_dashes": dashes_pairs,
            "rhetorical_style": (
                "engaging" if len(questions) + imperatives > len(sentences) * 0.2 else "declarative"
            ),
        }

    def calculate_metrics(
        self,
        text: str,
        base_patterns: dict[str, Any] | None = None,
        discovered_patterns: dict[str, list[str]] | None = None,
    ) -> dict[str, Any]:
        """Calculate comprehensive text metrics with deep linguistic analysis

        Args:
            text: Text to analyze
            base_patterns: Base linguistic patterns from YAML (optional)
            discovered_patterns: LLM-discovered patterns (optional)
        """
        sentences = self.split_sentences(text)
        paragraphs = self.split_paragraphs(text)
        words = self.tokenize_words(text)

        metrics = {
            "total_characters": len(text),
            "total_words": len(words),
            "total_sentences": len(sentences),
            "total_paragraphs": len(paragraphs),
            "avg_sentence_length": (len(words) / len(sentences) if sentences else 0.0),
            "avg_paragraph_length": (len(sentences) / len(paragraphs) if paragraphs else 0.0),
            "lexical_diversity": self.calculate_lexical_diversity(words),
        }

        # Detect primary language
        primary_lang, confidence = self.detect_language(text)
        metrics["primary_language"] = primary_lang
        metrics["language_confidence"] = confidence

        # Deep linguistic analysis
        metrics["voice_analysis"] = self.analyze_voice(text, primary_lang)
        metrics["transitions"] = self.extract_transition_words(
            text, primary_lang, base_patterns, discovered_patterns
        )
        metrics["sentence_complexity"] = self.analyze_sentence_complexity(text, primary_lang)
        metrics["rhetorical_devices"] = self.analyze_rhetorical_devices(text)

        # German-specific features
        if primary_lang == "de" or confidence > 0.3:
            metrics["german_features"] = self.detect_german_features(text, words)

        return metrics


class LLMAnalyzer:
    """Analyze text using local LLM"""

    def __init__(self, config: dict[str, Any], logger: logging.Logger):
        self.config = config
        self.logger = logger
        self.model = None
        self.tokenizer = None
        self.device = None
        self._load_model()

    def _load_model(self):
        """Load the LLM model based on configuration"""
        model_config = self.config["model"]
        model_type = model_config["type"]

        self.logger.info(f"Loading {model_type} model...")

        if model_type == "transformers":
            self._load_transformers_model(model_config)
        elif model_type == "llama-cpp":
            self._load_llama_cpp_model(model_config)
        else:
            raise ConfigurationError(f"Unknown model type: {model_type}")

    def _load_transformers_model(self, model_config: dict[str, Any]):
        """Load HuggingFace transformers model"""
        try:
            import torch
            from transformers import AutoModelForCausalLM, AutoTokenizer

            model_name = model_config["name"]
            device = model_config.get("device", "auto")

            self.logger.info(f"Loading model: {model_name}")

            # Load tokenizer
            self.tokenizer = AutoTokenizer.from_pretrained(model_name)

            # Load model
            if device == "auto":
                if torch.cuda.is_available():
                    self.device = "cuda"
                elif torch.backends.mps.is_available():
                    self.device = "mps"
                else:
                    self.device = "cpu"
            else:
                self.device = device

            self.logger.info(f"Using device: {self.device}")

            self.model = AutoModelForCausalLM.from_pretrained(
                model_name,
                torch_dtype=torch.float16 if self.device == "cuda" else torch.float32,
                device_map=self.device if self.device != "mps" else None,
            )

            if self.device == "mps":
                self.model = self.model.to(self.device)

            self.logger.info("Model loaded successfully")

        except Exception as e:
            raise ConfigurationError(f"Failed to load transformers model: {e}")

    def _load_llama_cpp_model(self, model_config: dict[str, Any]):
        """Load llama.cpp model (GGUF)"""
        try:
            from llama_cpp import Llama

            model_path = model_config.get("path")
            if not model_path:
                raise ConfigurationError("Model path required for llama-cpp")

            self.logger.info(f"Loading GGUF model: {model_path}")

            self.model = Llama(
                model_path=model_path,
                n_ctx=model_config.get("max_context_length", 4096),
                n_gpu_layers=-1 if model_config.get("device") == "cuda" else 0,
            )

            self.logger.info("Model loaded successfully")

        except Exception as e:
            raise ConfigurationError(f"Failed to load llama-cpp model: {e}")

    def _generate_transformers(self, prompt: str, max_tokens: int) -> str:
        """Generate text using transformers model"""
        import torch

        # Prepare input
        inputs = self.tokenizer(prompt, return_tensors="pt").to(self.device)

        # Generate
        with torch.no_grad():
            outputs = self.model.generate(
                **inputs,
                max_new_tokens=max_tokens,
                temperature=self.config["model"].get("temperature", 0.3),
                top_p=self.config["model"].get("top_p", 0.9),
                do_sample=True,
                pad_token_id=self.tokenizer.eos_token_id,
            )

        # Decode
        response = self.tokenizer.decode(outputs[0], skip_special_tokens=True)

        # Remove prompt from response
        if response.startswith(prompt):
            response = response[len(prompt) :].strip()

        return response

    def _generate_llama_cpp(self, prompt: str, max_tokens: int) -> str:
        """Generate text using llama-cpp model"""
        response = self.model(
            prompt,
            max_tokens=max_tokens,
            temperature=self.config["model"].get("temperature", 0.3),
            top_p=self.config["model"].get("top_p", 0.9),
        )

        return response["choices"][0]["text"].strip()

    def generate(self, prompt: str, max_tokens: int | None = None) -> str:
        """Generate text using the loaded model"""
        if max_tokens is None:
            max_tokens = self.config["model"].get("max_tokens", 2048)

        model_type = self.config["model"]["type"]

        try:
            if model_type == "transformers":
                return self._generate_transformers(prompt, max_tokens)
            elif model_type == "llama-cpp":
                return self._generate_llama_cpp(prompt, max_tokens)
        except Exception as e:
            raise AnalysisError(f"Generation failed: {e}")

    def _extract_json_from_response(self, response: str) -> dict[str, Any] | None:
        """Try multiple strategies to extract valid JSON from LLM response"""
        # Strategy 1: Look for JSON code blocks with ```json markers
        json_block_match = re.search(r"```json\s*(\{.*?\})\s*```", response, re.DOTALL)
        if json_block_match:
            try:
                return json.loads(json_block_match.group(1))
            except json.JSONDecodeError:
                pass

        # Strategy 2: Look for JSON code blocks with ``` markers (no language)
        code_block_match = re.search(r"```\s*(\{.*?\})\s*```", response, re.DOTALL)
        if code_block_match:
            try:
                return json.loads(code_block_match.group(1))
            except json.JSONDecodeError:
                pass

        # Strategy 3: Find the first complete JSON object (balanced braces)
        brace_count = 0
        start_idx = response.find("{")
        if start_idx != -1:
            for i in range(start_idx, len(response)):
                if response[i] == "{":
                    brace_count += 1
                elif response[i] == "}":
                    brace_count -= 1
                    if brace_count == 0:
                        # Found balanced JSON
                        try:
                            return json.loads(response[start_idx : i + 1])
                        except json.JSONDecodeError:
                            pass
                        break

        # Strategy 4: Try the greedy regex as last resort
        json_match = re.search(r"\{.*\}", response, re.DOTALL)
        if json_match:
            try:
                return json.loads(json_match.group())
            except json.JSONDecodeError:
                pass

        return None

    def analyze_style(
        self, text: str, metrics: dict[str, Any], language: str = "de", max_retries: int = 2
    ) -> dict[str, Any]:
        """Analyze writing style using LLM with robust JSON parsing and retry logic"""
        # Create analysis prompt with few-shot example
        if language == "de":
            system_prompt = """Du bist ein Experte für Schreibstilanalyse. Analysiere den folgenden Text und beschreibe den Schreibstil präzise.

WICHTIG: Antworte NUR mit validem JSON. Keine zusätzlichen Erklärungen vor oder nach dem JSON.

Beispiel für korrektes Format:
{
  "tone": "sachlich und akademisch mit neutralem Tonfall",
  "formality": "formal",
  "typical_elements": ["Fachterminologie", "Passivkonstruktionen", "Zitationen", "Strukturierte Argumentation"],
  "vocabulary_characteristics": "Wissenschaftlicher Wortschatz mit domänenspezifischen Fachbegriffen, präzise Formulierungen",
  "structural_patterns": ["Thesensätze am Anfang", "Evidenzbasierte Argumentation", "Absätze mit klarer Struktur"],
  "avoid": ["Umgangssprache", "persönliche Meinungen ohne Belege", "Kontraktionen"]
}

Achte beim Analysieren besonders auf:
- Tonfall (formell/informell, freundlich/sachlich, etc.)
- Sprachliche Besonderheiten (du/Sie, Fachbegriffe, Anglizismen)
- Typische Satzmuster und Strukturen
- Verwendung von Beispielen, Metaphern, Fragen
- Besondere stilistische Merkmale

Analysiere NUR die Prosa-Inhalte. Ignoriere Code, Formeln, Tabellen und Literaturverzeichnisse."""
        else:
            system_prompt = """You are an expert in writing style analysis. Analyze the following text and describe the writing style precisely.

IMPORTANT: Respond ONLY with valid JSON. No additional explanations before or after the JSON.

Example of correct format:
{
  "tone": "factual and academic with neutral tone",
  "formality": "formal",
  "typical_elements": ["technical terminology", "passive voice", "citations", "structured argumentation"],
  "vocabulary_characteristics": "Scientific vocabulary with domain-specific technical terms, precise formulations",
  "structural_patterns": ["topic sentences at start", "evidence-based reasoning", "clearly structured paragraphs"],
  "avoid": ["colloquialisms", "personal opinions without evidence", "contractions"]
}

Pay attention to:
- Tone (formal/informal, friendly/factual, etc.)
- Linguistic characteristics (technical terms, phrases)
- Typical sentence patterns and structures
- Use of examples, metaphors, questions
- Special stylistic features

Analyze ONLY the prose content. Ignore code, formulas, tables, and references."""

        # Try analysis with retry logic
        for attempt in range(max_retries + 1):
            try:
                # Sample text strategically (not just first 3000 chars)
                max_text_length = 8000 if attempt == 0 else 5000
                text_sample = self._sample_text_strategically(text, max_text_length)

                prompt = f"{system_prompt}\n\nText:\n{text_sample}\n\nJSON:"

                self.logger.info(
                    f"Generating style analysis (attempt {attempt + 1}/{max_retries + 1})..."
                )

                # Use lower temperature for JSON generation
                original_temp = self.config["model"].get("temperature", 0.3)
                self.config["model"]["temperature"] = 0.1  # Very low for structured output

                response = self.generate(prompt, max_tokens=1500)

                # Restore original temperature
                self.config["model"]["temperature"] = original_temp

                self.logger.debug(f"LLM response: {response[:200]}...")

                # Try to extract JSON
                style_data = self._extract_json_from_response(response)

                if style_data:
                    # Validate that we have the expected fields
                    required_fields = [
                        "tone",
                        "formality",
                        "typical_elements",
                        "vocabulary_characteristics",
                    ]
                    if all(field in style_data for field in required_fields):
                        self.logger.info("Successfully parsed style analysis")
                        return style_data
                    else:
                        self.logger.warning(f"JSON missing required fields: {required_fields}")

                # If we got here, parsing failed
                if attempt < max_retries:
                    self.logger.warning(
                        f"Failed to parse valid JSON, retrying ({attempt + 1}/{max_retries})..."
                    )
                    continue

            except Exception as e:
                self.logger.error(f"Analysis error on attempt {attempt + 1}: {e}")
                if attempt < max_retries:
                    continue

        # All retries failed - return fallback
        self.logger.error("All analysis attempts failed, returning fallback data")
        return {
            "tone": "Analysis failed - could not generate valid response",
            "formality": "unknown",
            "typical_elements": [],
            "vocabulary_characteristics": "Analysis failed",
            "structural_patterns": [],
            "avoid": [],
        }

    def _sample_text_strategically(self, text: str, max_length: int) -> str:
        """Sample text strategically rather than just taking the first N characters"""
        if len(text) <= max_length:
            return text

        # Split into paragraphs
        paragraphs = text.split("\n\n")

        # Try to get a good distribution: beginning, middle, end
        total_paras = len(paragraphs)
        if total_paras <= 3:
            return text[:max_length]

        # Sample from different sections
        samples = []
        char_count = 0

        # Beginning (first 2-3 paragraphs)
        for i in range(min(3, total_paras)):
            if char_count + len(paragraphs[i]) < max_length * 0.4:
                samples.append(paragraphs[i])
                char_count += len(paragraphs[i]) + 2

        # Middle (1-2 paragraphs from middle)
        mid_start = total_paras // 2
        for i in range(mid_start, min(mid_start + 2, total_paras)):
            if char_count + len(paragraphs[i]) < max_length * 0.7:
                samples.append(paragraphs[i])
                char_count += len(paragraphs[i]) + 2

        # End (last 1-2 paragraphs)
        for i in range(max(total_paras - 2, mid_start + 2), total_paras):
            if char_count + len(paragraphs[i]) < max_length:
                samples.append(paragraphs[i])
                char_count += len(paragraphs[i]) + 2

        result = "\n\n".join(samples)

        # Truncate if still too long
        if len(result) > max_length:
            result = result[:max_length] + "..."

        return result

    def extract_common_phrases(self, text: str, top_n: int = 20) -> list[str]:
        """Extract common phrases from text (2-3 word sequences)"""
        # German and English stopwords
        stopwords = {
            # German
            "der",
            "die",
            "das",
            "den",
            "dem",
            "des",
            "ein",
            "eine",
            "einer",
            "eines",
            "einem",
            "einen",
            "und",
            "oder",
            "aber",
            "auch",
            "wenn",
            "als",
            "wie",
            "bei",
            "mit",
            "nach",
            "von",
            "zu",
            "im",
            "am",
            "um",
            "an",
            "auf",
            "für",
            "ist",
            "sind",
            "war",
            "waren",
            "wird",
            "werden",
            "wurde",
            "wurden",
            "hat",
            "haben",
            "hatte",
            "hatten",
            "sich",
            "nicht",
            "nur",
            "noch",
            "mehr",
            "sehr",
            "dann",
            "dass",
            "dieser",
            "diese",
            "dieses",
            # English
            "the",
            "a",
            "an",
            "and",
            "or",
            "but",
            "if",
            "as",
            "at",
            "by",
            "for",
            "from",
            "in",
            "of",
            "on",
            "to",
            "with",
            "is",
            "are",
            "was",
            "were",
            "be",
            "been",
            "being",
            "have",
            "has",
            "had",
            "do",
            "does",
            "did",
            "will",
            "would",
            "can",
            "could",
            "should",
            "may",
            "might",
            "must",
            "this",
            "that",
            "these",
            "those",
            "it",
            "its",
            "not",
            "no",
            "yes",
            "so",
            "than",
            "then",
            "there",
            "when",
            "where",
            "which",
            "who",
            "why",
            "how",
        }

        words = re.findall(r"\b\w+\b", text.lower())

        # Filter out stopwords and numbers
        filtered_words = []
        for word in words:
            if word not in stopwords and not word.isdigit() and len(word) > 2:
                filtered_words.append(word)

        if len(filtered_words) < 3:
            return []

        # Extract bigrams and trigrams
        phrases = []

        # Bigrams
        for i in range(len(filtered_words) - 1):
            phrase = f"{filtered_words[i]} {filtered_words[i+1]}"
            # Skip if contains numbers or very short phrases
            if not any(char.isdigit() for char in phrase) and len(phrase) > 5:
                phrases.append(phrase)

        # Trigrams
        for i in range(len(filtered_words) - 2):
            phrase = f"{filtered_words[i]} {filtered_words[i+1]} {filtered_words[i+2]}"
            # Skip if contains numbers or very short phrases
            if not any(char.isdigit() for char in phrase) and len(phrase) > 8:
                phrases.append(phrase)

        # Count and return most common
        phrase_counts = Counter(phrases)
        return [phrase for phrase, count in phrase_counts.most_common(top_n)]


class WritingStyleAnalyzer:
    """Main analyzer orchestrator"""

    def __init__(self, config_path: str = "config.yaml"):
        self.config = load_config(config_path)
        self.logger = setup_logging(self.config)
        self.text_processor = TextProcessor(self.logger)
        self.llm_analyzer = LLMAnalyzer(self.config, self.logger)

    def collect_files(self, input_dir: Path) -> list[Path]:
        """Collect all text files from input directory"""
        file_config = self.config.get("files", {})
        extensions = file_config.get("extensions", [".txt", ".md"])
        recursive = file_config.get("recursive", True)

        files = []
        for ext in extensions:
            if recursive:
                files.extend(input_dir.rglob(f"*{ext}"))
            else:
                files.extend(input_dir.glob(f"*{ext}"))

        self.logger.info(f"Found {len(files)} files in {input_dir}")
        return files

    def read_file(self, file_path: Path) -> str:
        """Read file with proper encoding (supports .txt, .md, .tex, .pdf, .docx, .odt)"""
        try:
            suffix = file_path.suffix.lower()

            # Handle different file types
            if suffix == ".tex":
                return self._extract_latex_text(file_path)
            elif suffix == ".pdf":
                return self._extract_pdf_text(file_path)
            elif suffix == ".docx":
                return self._extract_docx_text(file_path)
            elif suffix == ".odt":
                return self._extract_odt_text(file_path)

            # Handle text files (.txt, .md)
            encoding = self.config.get("files", {}).get("encoding", "utf-8")
            with open(file_path, encoding=encoding) as f:
                return f.read()
        except Exception as e:
            self.logger.error(f"Failed to read {file_path}: {e}")
            return ""

    def _extract_latex_text(self, tex_path: Path) -> str:
        """Extract plain text from LaTeX file (.tex)"""
        try:
            from pylatexenc.latex2text import LatexNodes2Text

            # Read the .tex file
            encoding = self.config.get("files", {}).get("encoding", "utf-8")
            with open(tex_path, encoding=encoding) as f:
                latex_content = f.read()

            # Convert LaTeX to plain text
            # This handles:
            # - LaTeX commands (\textbf{}, \emph{}, etc.)
            # - Math mode ($...$, $$...$$)
            # - Special characters (ä, ö, ü, ß work natively!)
            # - Citations (\cite{})
            # - References (\ref{}, \label{})
            converter = LatexNodes2Text()
            plain_text = converter.latex_to_text(latex_content)

            self.logger.debug(f"Extracted {len(plain_text)} characters from LaTeX: {tex_path}")
            return plain_text

        except ImportError:
            self.logger.error("pylatexenc not installed. Install with: uv add pylatexenc")
            return ""
        except Exception as e:
            self.logger.error(f"Failed to extract text from LaTeX {tex_path}: {e}")
            return ""

    def _fix_umlaut_encoding(self, text: str) -> str:
        """Fix common German umlaut encoding issues from PDF extraction"""
        # Common patterns where umlauts get split into base + diacritic with space
        # ä = a + ¨ → "a " or "a¨"
        # ö = o + ¨ → "o " or "o¨"
        # ü = u + ¨ → "u " or "u¨"
        # ß often becomes "ß" or "ss"

        # Pattern 1: Base letter + space + combining character (most common)
        replacements = [
            (r"a\s+¨", "ä"),
            (r"o\s+¨", "ö"),
            (r"u\s+¨", "ü"),
            (r"A\s+¨", "Ä"),
            (r"O\s+¨", "Ö"),
            (r"U\s+¨", "Ü"),
            # Pattern 2: Base letter + space (contextual - only before consonants)
            (r"\ba\s+([bcdfghjklmnpqrstvwxyz])", r"ä\1"),
            (r"\bo\s+([bcdfghjklmnpqrstvwxyz])", r"ö\1"),
            (r"\bu\s+([bcdfghjklmnpqrstvwxyz])", r"ü\1"),
        ]

        fixed_text = text
        for pattern, replacement in replacements:
            fixed_text = re.sub(pattern, replacement, fixed_text, flags=re.IGNORECASE)

        return fixed_text

    def _extract_pdf_text(self, pdf_path: Path) -> str:
        """Extract text from PDF file using pdfplumber (better encoding handling)"""
        try:
            text_parts = []

            with pdfplumber.open(pdf_path) as pdf:
                for page_num, page in enumerate(pdf.pages, start=1):
                    try:
                        # pdfplumber provides better text extraction with proper encoding
                        page_text = page.extract_text()
                        if page_text:
                            # Fix any remaining umlaut encoding issues
                            page_text = self._fix_umlaut_encoding(page_text)
                            text_parts.append(page_text)
                        else:
                            self.logger.warning(
                                f"No text extracted from page {page_num} of {pdf_path}"
                            )
                    except Exception as e:
                        self.logger.warning(f"Error extracting page {page_num} of {pdf_path}: {e}")
                        continue

            full_text = "\n\n".join(text_parts)
            self.logger.debug(f"Extracted {len(full_text)} characters from PDF: {pdf_path}")
            return full_text

        except Exception as e:
            self.logger.error(f"Failed to extract text from PDF {pdf_path}: {e}")
            return ""

    def _extract_docx_text(self, docx_path: Path) -> str:
        """Extract text from DOCX file (Microsoft Word)"""
        try:
            doc = Document(docx_path)
            text_parts = []

            # Extract text from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)

            full_text = "\n\n".join(text_parts)
            self.logger.debug(f"Extracted {len(full_text)} characters from DOCX: {docx_path}")
            return full_text

        except Exception as e:
            self.logger.error(f"Failed to extract text from DOCX {docx_path}: {e}")
            return ""

    def _extract_odt_text(self, odt_path: Path) -> str:
        """Extract text from ODT file (LibreOffice Writer)"""
        try:
            doc = load_odt(str(odt_path))
            text_parts = []

            # Extract all text elements
            for element in doc.getElementsByType(text.P):
                paragraph_text = teletype.extractText(element)
                if paragraph_text.strip():
                    text_parts.append(paragraph_text)

            # Extract text from headers
            for element in doc.getElementsByType(text.H):
                header_text = teletype.extractText(element)
                if header_text.strip():
                    text_parts.append(header_text)

            full_text = "\n\n".join(text_parts)
            self.logger.debug(f"Extracted {len(full_text)} characters from ODT: {odt_path}")
            return full_text

        except Exception as e:
            self.logger.error(f"Failed to extract text from ODT {odt_path}: {e}")
            return ""

    def analyze_texts(self, input_dir: Path, profile_name: str) -> dict[str, Any]:
        """Analyze all texts and create profile"""
        files = self.collect_files(input_dir)

        if not files:
            raise AnalysisError(f"No text files found in {input_dir}")

        # Collect all texts
        self.logger.info("Reading files...")
        all_text = ""
        all_text_raw = ""  # Keep raw version for diagnostics
        language_counts = Counter()
        file_count = 0

        show_progress = self.config.get("logging", {}).get("progress_bars", True)
        file_iterator = tqdm(files, desc="Reading files") if show_progress else files

        for file_path in file_iterator:
            text = self.read_file(file_path)
            if text:
                all_text_raw += text + "\n\n"
                file_count += 1

                # Detect language
                lang, _ = self.text_processor.detect_language(text)
                language_counts[lang] += 1

        if not all_text_raw:
            raise AnalysisError("No text content found in files")

        self.logger.info(f"Read {file_count} files with {len(all_text_raw)} characters total")

        # Filter to prose content
        self.logger.info("Filtering to prose content (removing code, formulas, references)...")
        all_text = self.text_processor.filter_to_prose(all_text_raw)

        # Content quality diagnostics
        is_prose, prose_ratio = self.text_processor.is_likely_prose(all_text_raw)
        filtered_ratio = len(all_text) / len(all_text_raw) if len(all_text_raw) > 0 else 0

        self.logger.info(
            f"Prose filtering: kept {len(all_text)} chars ({filtered_ratio*100:.1f}% of original)"
        )
        self.logger.info(
            f"Content quality: {prose_ratio*100:.1f}% prose, {(1-prose_ratio)*100:.1f}% non-prose"
        )

        if filtered_ratio < 0.3:
            self.logger.warning(
                f"WARNING: Only {filtered_ratio*100:.1f}% of content is prose. "
                "Results may be inaccurate due to high code/formula/reference content."
            )

        if len(all_text) < 1000:
            self.logger.warning(
                f"WARNING: Only {len(all_text)} characters of prose found. "
                "Results may be inaccurate due to insufficient text."
            )

        if not all_text:
            raise AnalysisError(
                "No prose content found after filtering. Files contain only code/formulas/references."
            )

        self.logger.info(
            f"Analyzing {file_count} files with {len(all_text)} characters of prose..."
        )

        # Detect primary language first
        primary_lang, _ = self.text_processor.detect_language(all_text)

        # Load base linguistic patterns from YAML files
        self.logger.info(f"Loading base linguistic patterns for {primary_lang}...")
        base_patterns = load_linguistic_patterns(language=primary_lang, style="academic")

        # Discover additional patterns via LLM
        self.logger.info("Discovering transition words in your text via LLM...")
        discovered_patterns = discover_transition_words_llm(
            all_text, primary_lang, self.llm_analyzer, max_tokens=500
        )

        if discovered_patterns:
            self.logger.info(
                f"Discovered {sum(len(words) for words in discovered_patterns.values())} "
                "transition words from your text"
            )

        # Calculate metrics with hybrid patterns
        metrics = self.text_processor.calculate_metrics(
            all_text, base_patterns, discovered_patterns
        )

        # Analyze style with LLM (using filtered prose)
        style_data = self.llm_analyzer.analyze_style(all_text, metrics, primary_lang)

        # Extract common phrases from filtered prose
        common_phrases = self.llm_analyzer.extract_common_phrases(all_text)

        # Detect languages
        languages_detected = [lang for lang, count in language_counts.most_common() if count > 0]

        # Create profile with comprehensive analysis
        profile = {
            "profile_name": profile_name,
            "created_at": datetime.now().isoformat(),
            "analyzed_files": file_count,
            "primary_language": primary_lang,
            "languages_detected": languages_detected,
            "metrics": {
                "basic": {
                    "avg_sentence_length": round(metrics["avg_sentence_length"], 2),
                    "avg_paragraph_length": round(metrics["avg_paragraph_length"], 2),
                    "lexical_diversity": round(metrics["lexical_diversity"], 3),
                    "total_words": metrics["total_words"],
                    "total_sentences": metrics["total_sentences"],
                    "total_paragraphs": metrics["total_paragraphs"],
                },
                "voice_and_style": metrics["voice_analysis"],
                "sentence_structure": metrics["sentence_complexity"],
                "transitions": metrics["transitions"],
                "rhetorical": metrics["rhetorical_devices"],
            },
            "style_characteristics": {
                "tone": style_data.get("tone", "unknown"),
                "formality": style_data.get("formality", "unknown"),
                "typical_elements": style_data.get("typical_elements", []),
                "structural_patterns": style_data.get("structural_patterns", []),
            },
            "vocabulary": {
                "common_phrases": common_phrases[:10],
                "characteristics": style_data.get("vocabulary_characteristics", ""),
            },
            "avoid": style_data.get("avoid", []),
        }

        # Add German-specific features if detected
        if "german_features" in metrics:
            profile["german_features"] = metrics["german_features"]

        return profile

    def _generate_markdown_profile(self, profile: dict[str, Any]) -> str:
        """Generate markdown version of profile for AI agents"""
        md_lines = []

        # Header
        profile_name = profile.get("profile_name", "unknown")
        profile_type = profile.get("profile_type", "analyzed")
        primary_lang = profile.get("primary_language", "en")

        md_lines.append(f"# {profile_name.replace('-', ' ').title()} Writing Style Profile")
        md_lines.append("")
        md_lines.append(f"**Profile Name:** {profile_name}")
        md_lines.append(f"**Type:** {profile.get('description', profile_type)}")
        md_lines.append(
            f"**Language:** {'German (with Sie-form)' if primary_lang == 'de' else 'English'}"
        )
        md_lines.append("")
        md_lines.append("---")
        md_lines.append("")

        # Quick Instructions
        md_lines.append("## Quick Instructions")
        md_lines.append("")
        md_lines.append("Write in this style using these characteristics:")
        md_lines.append("")

        # Voice & Structure
        metrics = profile.get("metrics", {})
        basic = metrics.get("basic", {})
        voice = metrics.get("voice_and_style", {})

        md_lines.append("### Voice & Structure")
        if voice.get("passive_ratio"):
            passive_pct = int(voice["passive_ratio"] * 100)
            md_lines.append(f"- **Passive voice:** {passive_pct}%")
        if basic.get("avg_sentence_length"):
            md_lines.append(
                f"- **Sentence length:** ~{basic['avg_sentence_length']:.0f} words average"
            )
        if basic.get("avg_paragraph_length"):
            md_lines.append(
                f"- **Paragraph length:** ~{basic['avg_paragraph_length']:.1f} sentences per paragraph"
            )
        if basic.get("lexical_diversity"):
            md_lines.append(
                f"- **Lexical diversity:** {basic['lexical_diversity']:.3f} (use rich, varied vocabulary)"
            )

        sentence_struct = metrics.get("sentence_structure", {})
        if sentence_struct.get("complexity_ratio"):
            complexity_pct = int(sentence_struct["complexity_ratio"] * 100)
            md_lines.append(
                f"- **Complex sentences:** {complexity_pct}% (include subordinate clauses)"
            )
        md_lines.append("")

        # Transition Words
        transitions = metrics.get("transitions", {})
        by_category = transitions.get("by_category", {})

        if by_category:
            md_lines.append("### Transition Words")
            md_lines.append("")

            for category, data in by_category.items():
                if not isinstance(data, dict):
                    continue

                category_name = category.capitalize()
                count = data.get("count", 0)
                examples = data.get("examples", [])

                if count > 0:
                    md_lines.append(f"**{category_name}:**")
                    if examples:
                        examples_str = ", ".join(examples[:5])
                        md_lines.append(f"- Use: {examples_str}")
                    md_lines.append(f"- **Target:** ~{count} uses per document")
                    md_lines.append("")

        # Style Signature
        style_chars = profile.get("style_characteristics", {})
        if style_chars:
            md_lines.append("### Style Signature")
            if style_chars.get("tone"):
                md_lines.append(f"- **Tone:** {style_chars['tone']}")
            if sentence_struct.get("semicolons"):
                md_lines.append(
                    f"- **Semicolons:** Use ~{sentence_struct['semicolons']} per document"
                )
            if style_chars.get("formality"):
                md_lines.append(f"- **Formality:** {style_chars['formality']}")
            md_lines.append("")

        # What to Avoid
        avoid = profile.get("avoid", [])
        if avoid:
            md_lines.append("### What to Avoid")
            for item in avoid:
                md_lines.append(f"- {item}")
            md_lines.append("")

        md_lines.append("---")
        md_lines.append("")

        # Detailed Metrics
        md_lines.append("## Detailed Metrics")
        md_lines.append("")

        md_lines.append("### Basic Statistics")
        if basic.get("avg_sentence_length"):
            md_lines.append(f"- Average sentence length: {basic['avg_sentence_length']:.2f} words")
        if basic.get("avg_paragraph_length"):
            md_lines.append(
                f"- Average paragraph length: {basic['avg_paragraph_length']:.2f} sentences"
            )
        if basic.get("lexical_diversity"):
            md_lines.append(f"- Lexical diversity: {basic['lexical_diversity']:.3f}")
        if basic.get("total_words"):
            md_lines.append(f"- Total words analyzed: {basic['total_words']}")
        md_lines.append("")

        if voice:
            md_lines.append("### Voice Analysis")
            if voice.get("passive_ratio"):
                md_lines.append(f"- Passive sentence ratio: {voice['passive_ratio']*100:.1f}%")
            if voice.get("voice_style"):
                md_lines.append(f"- Voice style: {voice['voice_style']}")
            md_lines.append("")

        if by_category:
            md_lines.append("### Transition Word Breakdown")
            total = transitions.get("total_transitions", 0)
            density = transitions.get("transition_density", 0)
            if total > 0:
                md_lines.append(
                    f"- **Total transitions:** {total} (density: {density:.2f} per 100 words)"
                )

            for category, data in by_category.items():
                if isinstance(data, dict) and data.get("count", 0) > 0:
                    examples = ", ".join(data.get("examples", [])[:5])
                    md_lines.append(f"- **{category.capitalize()}:** {data['count']} ({examples})")
            md_lines.append("")

        # German-specific features
        german_features = profile.get("german_features", {})
        if german_features and primary_lang == "de":
            md_lines.append("## German-Specific Features")
            md_lines.append("")
            if german_features.get("formality"):
                md_lines.append(f"- **Formality:** {german_features['formality']}")
            if german_features.get("has_compound_words"):
                md_lines.append("- **Compound words:** Present")
                examples = german_features.get("compound_word_examples", [])
                if examples:
                    md_lines.append(f"- **Examples:** {', '.join(examples[:5])}")
            if german_features.get("uses_umlauts"):
                md_lines.append("- **Umlauts:** Proper ä, ö, ü, ß usage")
            md_lines.append("")

        # Footer
        md_lines.append("---")
        md_lines.append("")
        created = profile.get("created_at", "")
        if created:
            md_lines.append(f"**Created:** {created[:10]}")
        md_lines.append("**Source:** Generated by writing-style-analyzer")

        return "\n".join(md_lines)

    def save_profile(self, profile: dict[str, Any], output_path: Path):
        """Save profile to both JSON and Markdown files"""
        output_config = self.config.get("output", {})
        pretty = output_config.get("pretty_json", True)

        output_path.parent.mkdir(parents=True, exist_ok=True)

        # Save JSON
        with open(output_path, "w", encoding="utf-8") as f:
            if pretty:
                json.dump(profile, f, indent=2, ensure_ascii=False)
            else:
                json.dump(profile, f, ensure_ascii=False)

        self.logger.info(f"Profile saved to {output_path}")

        # Save Markdown
        md_path = output_path.with_suffix(".md")
        md_content = self._generate_markdown_profile(profile)

        with open(md_path, "w", encoding="utf-8") as f:
            f.write(md_content)

        self.logger.info(f"Markdown profile saved to {md_path}")


def main():
    """Main CLI entry point"""
    parser = argparse.ArgumentParser(
        description="Writing Style Analyzer - Analyze writing style using local LLMs"
    )
    parser.add_argument(
        "--input",
        "-i",
        type=str,
        required=True,
        help="Input directory containing text files",
    )
    parser.add_argument(
        "--output",
        "-o",
        type=str,
        required=True,
        help="Output path for profile JSON",
    )
    parser.add_argument(
        "--profile-type",
        "-t",
        type=str,
        default="general",
        help="Profile type (blog, social, roleplay, etc.)",
    )
    parser.add_argument(
        "--config",
        "-c",
        type=str,
        default="config.yaml",
        help="Path to config file (default: config.yaml)",
    )

    args = parser.parse_args()

    # Convert paths
    input_dir = Path(args.input)
    output_path = Path(args.output)

    # Validate input
    if not input_dir.exists():
        print(f"Error: Input directory does not exist: {input_dir}", file=sys.stderr)
        sys.exit(1)

    if not input_dir.is_dir():
        print(f"Error: Input path is not a directory: {input_dir}", file=sys.stderr)
        sys.exit(1)

    try:
        # Create analyzer
        analyzer = WritingStyleAnalyzer(args.config)

        # Analyze
        profile_name = args.profile_type
        profile = analyzer.analyze_texts(input_dir, profile_name)

        # Save
        analyzer.save_profile(profile, output_path)

        print("\n✓ Analysis complete!")
        print(f"  Profile: {profile_name}")
        print(f"  Files analyzed: {profile['analyzed_files']}")
        print(f"  Primary language: {profile['primary_language']}")
        print(f"  Output: {output_path}")

    except (ConfigurationError, AnalysisError) as e:
        print(f"Error: {e}", file=sys.stderr)
        sys.exit(1)
    except KeyboardInterrupt:
        print("\nInterrupted by user", file=sys.stderr)
        sys.exit(130)
    except Exception as e:
        print(f"Unexpected error: {e}", file=sys.stderr)
        import traceback

        traceback.print_exc()
        sys.exit(1)


if __name__ == "__main__":
    main()
